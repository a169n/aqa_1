from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


REPO_ROOT = Path(__file__).resolve().parents[2]
SOURCE_MD = REPO_ROOT / "docs" / "qa" / "assignment-2-report.md"
OUTPUT_DOCX = REPO_ROOT / "output" / "doc" / "Assignment_2_Report.docx"
SCREENSHOTS_DIR = REPO_ROOT / "docs" / "qa" / "screenshots"

IMAGE_BY_PLACEHOLDER = {
    "pipeline-overview.png": SCREENSHOTS_DIR / "pipeline-overview.png",
    "workflow-pass.png": SCREENSHOTS_DIR / "workflow-pass.png",
    "qa-summary.png": SCREENSHOTS_DIR / "qa-summary.png",
    "coverage-summary.png": SCREENSHOTS_DIR / "coverage-summary.png",
    "playwright-report.png": SCREENSHOTS_DIR / "playwright-report.png",
    "automation-coverage-chart.png": SCREENSHOTS_DIR / "automation-coverage-chart.png",
    "execution-time-chart.png": SCREENSHOTS_DIR / "execution-time-chart.png",
}

INLINE_TOKEN_RE = re.compile(r"(\*\*.+?\*\*|`.+?`)")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
ORDERED_ITEM_RE = re.compile(r"^\d+\.\s+")
FIGURE_RE = re.compile(r"^\*\*Figure Placeholder (\d+)\.\s+(.+?)\*\*$")
INSERT_FIGURE_RE = re.compile(r"^\[Insert Figure:\s*(.+?)\]$")


def is_path_like(value: str) -> bool:
    if "/" in value or "\\" in value:
        return True
    if re.search(r"\b[\w.-]+\.(?:png|jpg|jpeg|svg|json|md|yml|yaml|ts|tsx|js|cjs|spec|py|docx|pdf)\b", value):
        return True
    if value.startswith("http://") or value.startswith("https://"):
        return True
    return False


def normalize_source_text(text: str) -> str:
    replacements = {
        "(1)what": "(1) what",
        "(2)why": "(2) why",
        "(3)how": "(3) how",
        "(4)what": "(4) what",
        "refresh rotation,`/auth/me`": "refresh rotation, `/auth/me`",
        "for Figure Placeholders `6` and `7`": "for the automation-coverage and execution-time figures",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text


def configure_page(section) -> None:
    section.page_width = Inches(11)
    section.page_height = Inches(17)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)


def apply_run_font(run, *, bold: bool = False, italic: bool = False) -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic


def add_text_runs(paragraph, text: str, *, paragraph_bold: bool = False) -> None:
    parts = INLINE_TOKEN_RE.split(text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2].replace("`", ""))
            apply_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            token = part[1:-1]
            run = paragraph.add_run(token)
            apply_run_font(run, bold=paragraph_bold, italic=is_path_like(token))
        else:
            run = paragraph.add_run(part)
            apply_run_font(run, bold=paragraph_bold)


def base_paragraph(document: Document, *, align: WD_ALIGN_PARAGRAPH | None = None):
    paragraph = document.add_paragraph()
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.15
    if align is not None:
        paragraph.alignment = align
    return paragraph


def add_paragraph(document: Document, text: str, *, bold: bool = False, align=None) -> None:
    paragraph = base_paragraph(document, align=align)
    add_text_runs(paragraph, text, paragraph_bold=bold)


def split_heading_text(text: str) -> tuple[str | None, str]:
    match = re.match(r"^(\d+(?:\.\d+)*)\s+(.*)$", text)
    if match:
        return match.group(1), match.group(2)
    return None, text


def add_heading(document: Document, text: str, level: int) -> None:
    align = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    paragraph = base_paragraph(document, align=align)
    paragraph.paragraph_format.space_before = Pt(8 if level <= 2 else 4)
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    apply_run_font(run, bold=True)


def clean_cell_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip())


def is_markdown_separator_row(cells: list[str]) -> bool:
    normalized = [cell.strip() for cell in cells]
    if not normalized:
        return False
    return all(set(cell) <= {":", "-"} and "-" in cell for cell in normalized if cell) and all(
        cell for cell in normalized
    )


def set_cell_border(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)

    for edge in ("top", "left", "bottom", "right"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def build_table(document: Document, rows: list[list[str]], caption: str) -> None:
    add_paragraph(document, caption, bold=True)
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_idx, row in enumerate(rows):
        for col_idx, cell_text in enumerate(row):
            cell = table.cell(row_idx, col_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.0
            add_text_runs(paragraph, clean_cell_text(cell_text), paragraph_bold=row_idx == 0)
            set_cell_border(cell)

    document.add_paragraph()


def insert_figure(document: Document, number: str, title: str, image_name: str) -> None:
    image_path = IMAGE_BY_PLACEHOLDER.get(image_name)
    if image_path is None or not image_path.exists():
        raise FileNotFoundError(f"Missing figure asset: {image_name}")

    add_paragraph(document, f"Figure {number}. {title}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    paragraph = base_paragraph(document, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(8.9))
    paragraph.paragraph_format.space_after = Pt(10)


def extract_figure_name(line: str) -> str | None:
    cleaned = line.strip()
    if cleaned.startswith("`") and cleaned.endswith("`"):
        cleaned = cleaned[1:-1].strip()
    match = INSERT_FIGURE_RE.match(cleaned)
    if not match:
        return None
    content = match.group(1).strip()
    if " - " in content:
        return content.split(" - ", 1)[0].strip()
    return content


def build_docx() -> Path:
    text = normalize_source_text(SOURCE_MD.read_text(encoding="utf-8"))
    lines = text.splitlines()

    document = Document()
    configure_page(document.sections[0])

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal_style.font.size = Pt(12)

    table_counts: dict[str, int] = defaultdict(int)
    current_heading_text = ""
    paragraph_buffer: list[str] = []

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_paragraph(document, " ".join(part.strip() for part in paragraph_buffer if part.strip()))
            paragraph_buffer = []

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph()
            i += 1
            continue

        figure_match = FIGURE_RE.match(stripped)
        if figure_match:
            flush_paragraph()
            figure_number = figure_match.group(1)
            figure_title = figure_match.group(2)

            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines):
                image_name = extract_figure_name(lines[j])
                if image_name:
                    insert_figure(document, figure_number, figure_title, image_name)
                    i = j + 1
                    continue
            raise ValueError(f"Could not resolve image placeholder for figure: {figure_title}")

        if INSERT_FIGURE_RE.match(stripped):
            i += 1
            continue

        heading_match = HEADING_RE.match(stripped)
        if heading_match:
            flush_paragraph()
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            add_heading(document, heading_text, level)
            current_heading_text = heading_text
            i += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            parsed_rows = [[cell.strip() for cell in row.strip("|").split("|")] for row in table_lines]
            rows = [row for row in parsed_rows if not is_markdown_separator_row(row)]

            heading_number, heading_title = split_heading_text(current_heading_text)
            heading_key = current_heading_text or "Table"
            table_counts[heading_key] += 1
            table_idx = table_counts[heading_key]
            if heading_number:
                if table_idx == 1:
                    caption = f"Table {heading_number}. {heading_title}"
                else:
                    caption = f"Table {heading_number}.{table_idx}. {heading_title}"
            else:
                caption = f"Table. {heading_title if current_heading_text else 'Data'}"

            build_table(document, rows, caption)
            continue

        if stripped.startswith("- ") or ORDERED_ITEM_RE.match(stripped):
            flush_paragraph()
            add_paragraph(document, stripped)
            i += 1
            continue

        if re.match(r"^\*\*[^*]+:\*\*", stripped):
            flush_paragraph()
            add_paragraph(document, stripped)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph()

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)
    return OUTPUT_DOCX


def iter_all_paragraphs(document: Document) -> Iterable:
    for paragraph in document.paragraphs:
        yield paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph


def validate_docx(docx_path: Path) -> None:
    document = Document(str(docx_path))
    section = document.sections[0]

    if round(section.page_width.inches, 2) != 11.00 or round(section.page_height.inches, 2) != 17.00:
        raise ValueError("Document page size is not tabloid (11 x 17 inches).")

    all_text = "\n".join(paragraph.text for paragraph in iter_all_paragraphs(document))
    forbidden = ["Insert Figure", "Figure Placeholder", "manually prepared"]
    for token in forbidden:
        if token in all_text:
            raise ValueError(f"Found leftover placeholder text: {token}")

    if len(document.inline_shapes) != 7:
        raise ValueError(f"Expected 7 inserted figures, found {len(document.inline_shapes)}.")

    for paragraph in iter_all_paragraphs(document):
        for run in paragraph.runs:
            if run.text and run.font.name not in (None, "Times New Roman"):
                raise ValueError(f"Unexpected font detected: {run.font.name}")
            if run.text and run.font.size not in (None, Pt(12)):
                raise ValueError("Unexpected font size detected.")


if __name__ == "__main__":
    output = build_docx()
    validate_docx(output)
    print(f"Generated {output}")
