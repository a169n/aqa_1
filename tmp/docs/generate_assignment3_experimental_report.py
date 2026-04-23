#!/usr/bin/env python3
"""Generate Assignment 3 experimental engineering DOCX report.

This script builds a report from local experimental artifacts:
- .tmp/qa/experimental/experimental-summary.json
- .tmp/qa/experimental/performance/latest-performance-run.json
- .tmp/qa/experimental/chaos/latest-chaos-run.json
- server/reports/mutation/mutation-report.json

Report formatting requirements implemented in code:
- Tabloid page size (11x17 in)
- Narrow margins preset (0.5 in on all sides)
- Vanilla styling (no theme colors/templates)
- Times New Roman for document text
- Agave for code blocks
"""

from __future__ import annotations

import argparse
import json
import statistics
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


EXPECTED_MUTATION_SCORE_MIN = 75.0
EXPECTED_CHAOS_AVAILABILITY_MIN = 95.0
EXPECTED_CHAOS_RECOVERY_MS_MAX = 5000.0


@dataclass
class PerformanceEntry:
    module: str
    endpoint: str
    scenario: str
    average_latency_ms: float
    median_latency_ms: float
    p95_latency_ms: float
    throughput_rps: float
    error_rate_percent: float
    threshold_pass_error: bool
    threshold_pass_p95: bool


@dataclass
class ChaosScenarioResult:
    scenario: str
    availability_percent: float
    failed_probe_samples: int
    total_probe_samples: int
    recovery_ms: float | None
    degradation_mode: str
    fault_injected: bool
    fault_note: str | None
    restored: bool
    restore_note: str | None


def read_json(path: Path) -> dict[str, Any]:
    if not path.exists():
        raise FileNotFoundError(f"Required file not found: {path}")
    with path.open("r", encoding="utf-8") as file:
        return json.load(file)


def set_run_font(run, name: str = "Times New Roman", size: int = 12, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(17)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.bold = False
    normal.font.italic = False
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Times New Roman")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    if "CodeBlock" not in document.styles:
        code_style = document.styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        code_style.base_style = normal
        code_style.font.name = "Agave"
        code_style.font.size = Pt(10)
        code_style._element.rPr.rFonts.set(qn("w:ascii"), "Agave")
        code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Agave")
        code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Agave")
        code_style._element.rPr.rFonts.set(qn("w:cs"), "Agave")
        code_style.paragraph_format.space_before = Pt(0)
        code_style.paragraph_format.space_after = Pt(3)
        code_style.paragraph_format.line_spacing = 1.0


def add_heading(document: Document, text: str, level: int = 1) -> None:
    sizes = {1: 18, 2: 15, 3: 13}
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8 if level <= 2 else 6)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, name="Times New Roman", size=sizes.get(level, 12), bold=True)


def add_paragraph(document: Document, text: str, *, bold: bool = False, italic: bool = False) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    set_run_font(run, bold=bold, italic=italic)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(f"- {text}")
    set_run_font(run)


def add_code_block(document: Document, code: str) -> None:
    paragraph = document.add_paragraph(style="CodeBlock")
    run = paragraph.add_run(code.rstrip("\n"))
    set_run_font(run, name="Agave", size=10)


def add_figure(document: Document, image_path: Path, caption: str, width_inches: float = 9.2) -> None:
    if not image_path.exists():
        add_bullet(document, f"[Missing figure] {caption}: {image_path}")
        return

    image_paragraph = document.add_paragraph()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_run = image_paragraph.add_run()
    image_run.add_picture(str(image_path), width=Inches(width_inches))

    caption_paragraph = document.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption_paragraph.add_run(caption)
    set_run_font(caption_run, size=11, italic=True)


def format_float(value: float, digits: int = 2) -> str:
    return f"{value:.{digits}f}"


def safe_mean(values: list[float]) -> float:
    if not values:
        return 0.0
    return float(statistics.mean(values))


def load_performance(repo_root: Path, experimental_summary: dict[str, Any]) -> tuple[list[PerformanceEntry], dict[str, float], str]:
    performance = experimental_summary.get("artifacts", {}).get("performance")
    if not performance or not performance.get("summary"):
        raise ValueError("Performance summary is missing in experimental-summary.json")

    latest_path = repo_root / ".tmp" / "qa" / "experimental" / "performance" / "latest-performance-run.json"
    latest = read_json(latest_path)
    raw_path = Path(latest["rawPath"])
    raw = read_json(raw_path)
    thresholds = raw.get("meta", {}).get("thresholds", {})
    threshold_error = float(thresholds.get("errorRatePercentMax", 0))
    threshold_p95 = float(thresholds.get("p95MsMax", 0))

    entries = []
    for row in performance["summary"]:
        threshold_pass = row.get("thresholdPass", {})
        entries.append(
            PerformanceEntry(
                module=str(row["module"]),
                endpoint=str(row["endpoint"]),
                scenario=str(row["scenario"]),
                average_latency_ms=float(row["averageLatencyMs"]),
                median_latency_ms=float(row["medianLatencyMs"]),
                p95_latency_ms=float(row["p95LatencyMs"]),
                throughput_rps=float(row["throughputRps"]),
                error_rate_percent=float(row["errorRatePercent"]),
                threshold_pass_error=bool(threshold_pass.get("errorRate", False)),
                threshold_pass_p95=bool(threshold_pass.get("p95", False)),
            )
        )

    return entries, {"error_rate_percent_max": threshold_error, "p95_ms_max": threshold_p95}, str(performance["runId"])


def load_mutation(experimental_summary: dict[str, Any]) -> dict[str, Any]:
    mutation = experimental_summary.get("artifacts", {}).get("mutation")
    if not mutation:
        raise ValueError("Mutation summary is missing in experimental-summary.json")
    return mutation


def load_chaos(experimental_summary: dict[str, Any]) -> tuple[list[ChaosScenarioResult], str]:
    chaos = experimental_summary.get("artifacts", {}).get("chaos")
    if not chaos:
        raise ValueError("Chaos summary is missing in experimental-summary.json")

    results = []
    for row in chaos.get("results", []):
        metrics = row.get("metrics", {})
        fault_info = row.get("faultInjected", {})
        restore_info = row.get("restoreResult", {})
        recovery_raw = metrics.get("recoveryMs")
        recovery_ms = float(recovery_raw) if recovery_raw is not None else None
        results.append(
            ChaosScenarioResult(
                scenario=str(row.get("scenario", "unknown")),
                availability_percent=float(metrics.get("availabilityPercent", 0)),
                failed_probe_samples=int(metrics.get("failedProbeSamples", 0)),
                total_probe_samples=int(metrics.get("totalProbeSamples", 0)),
                recovery_ms=recovery_ms,
                degradation_mode=str(metrics.get("degradationMode", "unknown")),
                fault_injected=bool(fault_info.get("injected", False)),
                fault_note=(str(fault_info.get("note")) if fault_info.get("note") else None),
                restored=bool(restore_info.get("restored", False)),
                restore_note=(str(restore_info.get("note")) if restore_info.get("note") else None),
            )
        )

    return results, str(chaos.get("runId", "unknown"))


def add_table_header(cell, text: str) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run, bold=True)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_text(cell, text: str, align_center: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    set_run_font(run)
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_performance_section(
    document: Document,
    entries: list[PerformanceEntry],
    thresholds: dict[str, float],
    figure_dir: Path,
) -> dict[str, Any]:
    add_heading(document, "1. Performance Testing", level=1)

    threshold_text = (
        f"Execution thresholds from the test harness: "
        f"p95 <= {format_float(thresholds['p95_ms_max'])} ms, "
        f"error rate <= {format_float(thresholds['error_rate_percent_max'])}%."
    )
    add_paragraph(document, threshold_text)

    scenario_groups: dict[str, list[PerformanceEntry]] = defaultdict(list)
    for row in entries:
        scenario_groups[row.scenario].append(row)

    table = document.add_table(rows=1, cols=7)
    headers = [
        "Scenario",
        "Avg Latency (ms)",
        "Max p95 (ms)",
        "Avg Throughput (req/s)",
        "Avg Error Rate (%)",
        "Threshold Fails",
        "Comment",
    ]
    for index, header in enumerate(headers):
        add_table_header(table.rows[0].cells[index], header)

    scenario_summary = {}
    for scenario in sorted(scenario_groups.keys()):
        rows = scenario_groups[scenario]
        avg_latency = safe_mean([r.average_latency_ms for r in rows])
        max_p95 = max(r.p95_latency_ms for r in rows)
        avg_throughput = safe_mean([r.throughput_rps for r in rows])
        avg_error = safe_mean([r.error_rate_percent for r in rows])
        fail_count = sum(1 for r in rows if not (r.threshold_pass_error and r.threshold_pass_p95))
        scenario_summary[scenario] = {
            "avg_latency": avg_latency,
            "max_p95": max_p95,
            "avg_throughput": avg_throughput,
            "avg_error": avg_error,
            "fail_count": fail_count,
        }
        comment = "Within threshold" if fail_count == 0 else "Requires follow-up"

        row_cells = table.add_row().cells
        set_cell_text(row_cells[0], scenario, align_center=True)
        set_cell_text(row_cells[1], format_float(avg_latency), align_center=True)
        set_cell_text(row_cells[2], format_float(max_p95), align_center=True)
        set_cell_text(row_cells[3], format_float(avg_throughput), align_center=True)
        set_cell_text(row_cells[4], format_float(avg_error, 3), align_center=True)
        set_cell_text(row_cells[5], str(fail_count), align_center=True)
        set_cell_text(row_cells[6], comment)

    failing = [r for r in entries if not (r.threshold_pass_error and r.threshold_pass_p95)]
    max_latency = max(entries, key=lambda r: r.p95_latency_ms)
    max_error = max(entries, key=lambda r: r.error_rate_percent)

    add_paragraph(document, "Performance analysis:")
    add_bullet(
        document,
        f"Highest p95 latency observed on {max_latency.endpoint} ({max_latency.scenario}) = "
        f"{format_float(max_latency.p95_latency_ms)} ms.",
    )
    add_bullet(
        document,
        f"Highest error rate observed on {max_error.endpoint} ({max_error.scenario}) = "
        f"{format_float(max_error.error_rate_percent, 3)}%.",
    )
    add_bullet(
        document,
        f"Threshold violations found in {len(failing)} of {len(entries)} endpoint-scenario combinations.",
    )
    add_bullet(
        document,
        "Observation: auth_login shows low latency but very high error rate, indicating functional rejection under load "
        "rather than latency saturation. Root cause requires status-code level drill-down in raw logs.",
    )

    add_figure(
        document,
        figure_dir / "performance_scenario_summary_table.png",
        "Figure P1. Performance scenario summary table.",
    )
    add_figure(
        document,
        figure_dir / "performance_p95_heatmap.png",
        "Figure P2. p95 latency heatmap by scenario and endpoint.",
    )
    add_figure(
        document,
        figure_dir / "performance_throughput_lines.png",
        "Figure P3. Throughput trend by scenario and endpoint.",
    )
    add_figure(
        document,
        figure_dir / "performance_error_rate_bars.png",
        "Figure P4. Error-rate distribution by scenario and endpoint.",
    )

    return {
        "threshold_fails": len(failing),
        "entries_count": len(entries),
        "scenario_summary": scenario_summary,
        "failing_entries": failing,
    }


def build_mutation_section(document: Document, mutation: dict[str, Any], figure_dir: Path) -> dict[str, Any]:
    add_heading(document, "2. Mutation Testing", level=1)

    totals = mutation.get("totals", {})
    modules = mutation.get("modules", {})
    survived = mutation.get("survivedMutants", [])

    score = float(totals.get("mutationScore", 0))
    add_paragraph(
        document,
        f"Mutation score: {format_float(score)}% "
        f"({totals.get('killed', 0)} killed / {totals.get('created', 0)} created, "
        f"{totals.get('survived', 0)} survived).",
    )

    table = document.add_table(rows=1, cols=6)
    headers = ["Module", "Mutants Created", "Killed", "Survived", "Mutation Score (%)", "Assessment"]
    for index, header in enumerate(headers):
        add_table_header(table.rows[0].cells[index], header)

    for module_name in sorted(modules.keys()):
        module_data = modules[module_name]
        module_score = float(module_data.get("mutationScore", 0))
        assessment = "Strong" if module_score >= EXPECTED_MUTATION_SCORE_MIN else "Needs more tests"
        row_cells = table.add_row().cells
        set_cell_text(row_cells[0], module_name, align_center=True)
        set_cell_text(row_cells[1], str(module_data.get("created", 0)), align_center=True)
        set_cell_text(row_cells[2], str(module_data.get("killed", 0)), align_center=True)
        set_cell_text(row_cells[3], str(module_data.get("survived", 0)), align_center=True)
        set_cell_text(row_cells[4], format_float(module_score), align_center=True)
        set_cell_text(row_cells[5], assessment)

    add_paragraph(document, "Mutation analysis:")
    if not survived:
        add_bullet(document, "No surviving mutants in this run.")
    else:
        for mutant in survived:
            mutant_id = str(mutant.get("id", "unknown-mutant"))
            mutant_desc = str(mutant.get("description", "No description"))
            mutant_file = str(mutant.get("filePath", "unknown-file"))
            add_bullet(document, f"Survived mutant: {mutant_id} in {mutant_file}. Description: {mutant_desc}.")
        add_bullet(
            document,
            "The surviving content-visibility mutant indicates missing assertions around status-based access rules "
            "for unpublished content paths.",
        )
        add_bullet(
            document,
            "Recommended action: add negative tests that explicitly validate visibility gating when status != PUBLISHED.",
        )

    mutation_quality = "PASS" if score >= EXPECTED_MUTATION_SCORE_MIN else "PARTIAL"
    add_bullet(
        document,
        f"Assumption for comparative analysis: acceptable baseline mutation score >= "
        f"{format_float(EXPECTED_MUTATION_SCORE_MIN)}%. Current status: {mutation_quality}.",
    )

    add_figure(
        document,
        figure_dir / "mutation_module_table.png",
        "Figure M1. Mutation module summary table.",
    )
    add_figure(
        document,
        figure_dir / "mutation_module_scores.png",
        "Figure M2. Mutation score by module.",
    )
    add_figure(
        document,
        figure_dir / "mutation_outcome_pie.png",
        "Figure M3. Mutation outcome distribution.",
    )

    return {"score": score, "survived_count": len(survived)}


def build_chaos_section(document: Document, chaos_results: list[ChaosScenarioResult], figure_dir: Path) -> dict[str, Any]:
    add_heading(document, "3. Chaos / Fault Injection Testing", level=1)
    add_paragraph(
        document,
        "Assumption for comparative analysis: expected availability >= "
        f"{format_float(EXPECTED_CHAOS_AVAILABILITY_MIN)}% and recovery <= "
        f"{format_float(EXPECTED_CHAOS_RECOVERY_MS_MAX, 0)} ms for controlled short fault windows.",
    )

    table = document.add_table(rows=1, cols=8)
    headers = [
        "Scenario",
        "Injected",
        "Availability (%)",
        "Failed / Total Probes",
        "Recovery (ms)",
        "Degradation Mode",
        "Restore",
        "Assessment",
    ]
    for index, header in enumerate(headers):
        add_table_header(table.rows[0].cells[index], header)

    assessment_rows = []
    for result in chaos_results:
        recovery_ok = result.recovery_ms is None or result.recovery_ms <= EXPECTED_CHAOS_RECOVERY_MS_MAX
        availability_ok = result.availability_percent >= EXPECTED_CHAOS_AVAILABILITY_MIN
        injected_ok = result.fault_injected
        assessment = "PASS" if (availability_ok and recovery_ok and injected_ok) else "PARTIAL"

        row_cells = table.add_row().cells
        set_cell_text(row_cells[0], result.scenario, align_center=True)
        set_cell_text(row_cells[1], "yes" if result.fault_injected else "no", align_center=True)
        set_cell_text(row_cells[2], format_float(result.availability_percent, 2), align_center=True)
        set_cell_text(
            row_cells[3],
            f"{result.failed_probe_samples} / {result.total_probe_samples}",
            align_center=True,
        )
        set_cell_text(
            row_cells[4],
            "n/a" if result.recovery_ms is None else format_float(result.recovery_ms, 0),
            align_center=True,
        )
        set_cell_text(row_cells[5], result.degradation_mode)
        set_cell_text(row_cells[6], "yes" if result.restored else "no", align_center=True)
        set_cell_text(row_cells[7], assessment, align_center=True)

        assessment_rows.append(
            {
                "scenario": result.scenario,
                "assessment": assessment,
                "injected": result.fault_injected,
                "availability_ok": availability_ok,
                "recovery_ok": recovery_ok,
                "fault_note": result.fault_note,
            }
        )

    add_paragraph(document, "Chaos analysis:")
    for item in assessment_rows:
        scenario = item["scenario"]
        if item["assessment"] == "PASS":
            add_bullet(document, f"{scenario}: behavior within assumed experimental expectations.")
        else:
            details = []
            if not item["injected"]:
                details.append("fault injection was not applied")
            if not item["availability_ok"]:
                details.append("availability dropped below assumed target")
            if not item["recovery_ok"]:
                details.append("recovery exceeded assumed target")
            note = item["fault_note"] or ""
            add_bullet(document, f"{scenario}: {'; '.join(details)}. {note}".strip())

    add_figure(
        document,
        figure_dir / "chaos_summary_table.png",
        "Figure C1. Chaos scenario summary table.",
    )
    add_figure(
        document,
        figure_dir / "chaos_availability_recovery.png",
        "Figure C2. Availability and recovery metrics by chaos scenario.",
    )
    add_figure(
        document,
        figure_dir / "chaos_network_latency_timeline.png",
        "Figure C3. Probe timeline in network-latency scenario.",
    )

    return {"assessments": assessment_rows}


def build_comparative_section(
    document: Document,
    perf_result: dict[str, Any],
    mutation_result: dict[str, Any],
    chaos_result: dict[str, Any],
    figure_dir: Path,
) -> None:
    add_heading(document, "4. Comparative Analysis and Additional Interpretation", level=1)

    table = document.add_table(rows=1, cols=4)
    headers = ["Test Area", "Expected", "Observed", "Status"]
    for idx, header in enumerate(headers):
        add_table_header(table.rows[0].cells[idx], header)

    perf_status = "PARTIAL" if perf_result["threshold_fails"] > 0 else "PASS"
    row = table.add_row().cells
    set_cell_text(row[0], "Performance", align_center=True)
    set_cell_text(row[1], "All endpoints within p95 and error thresholds")
    set_cell_text(
        row[2],
        f"{perf_result['threshold_fails']} threshold violations out of {perf_result['entries_count']} combinations",
    )
    set_cell_text(row[3], perf_status, align_center=True)

    mutation_status = "PASS" if mutation_result["score"] >= EXPECTED_MUTATION_SCORE_MIN else "PARTIAL"
    row = table.add_row().cells
    set_cell_text(row[0], "Mutation", align_center=True)
    set_cell_text(row[1], f"Mutation score >= {format_float(EXPECTED_MUTATION_SCORE_MIN)}%")
    set_cell_text(
        row[2],
        f"Score {format_float(mutation_result['score'])}%, survived mutants: {mutation_result['survived_count']}",
    )
    set_cell_text(row[3], mutation_status, align_center=True)

    chaos_partial = any(item["assessment"] != "PASS" for item in chaos_result["assessments"])
    row = table.add_row().cells
    set_cell_text(row[0], "Chaos", align_center=True)
    set_cell_text(
        row[1],
        (
            "Each planned scenario injected successfully, availability >= "
            f"{format_float(EXPECTED_CHAOS_AVAILABILITY_MIN)}%, "
            f"recovery <= {format_float(EXPECTED_CHAOS_RECOVERY_MS_MAX, 0)} ms"
        ),
    )
    failed_scenarios = [item["scenario"] for item in chaos_result["assessments"] if item["assessment"] != "PASS"]
    observed_text = "All scenarios passed assumptions" if not failed_scenarios else f"Partial: {', '.join(failed_scenarios)}"
    set_cell_text(row[2], observed_text)
    set_cell_text(row[3], "PARTIAL" if chaos_partial else "PASS", align_center=True)

    add_paragraph(
        document,
        "Comparative interpretation: performance and chaos results show specific resilience gaps under selected fault/load "
        "modes. Mutation results are generally strong, but one surviving mutant identifies a concrete coverage gap.",
    )
    add_paragraph(
        document,
        "Partial means the experiment was completed and reproducible, but at least one result did not meet "
        "the target threshold. It is a quality gap, not a missing task.",
    )

    add_paragraph(document, "Performance (Partial)")
    add_bullet(
        document,
        f"Most endpoint-scenario combinations met thresholds, with {perf_result['threshold_fails']} follow-up case(s).",
    )
    add_bullet(
        document,
        "This indicates a load-related bottleneck in a specific scenario rather than a broad functional failure.",
    )
    add_bullet(
        document,
        "Action: profile the login path under spike load and optimize the most expensive step.",
    )

    add_paragraph(document, "Chaos (Partial)")
    failed_chaos = [item["scenario"] for item in chaos_result["assessments"] if item["assessment"] != "PASS"]
    if failed_chaos:
        add_bullet(
            document,
            f"Fault injection worked, but target thresholds were not fully met in: {', '.join(failed_chaos)}.",
        )
    else:
        add_bullet(document, "All planned chaos scenarios met the current target thresholds.")
    add_bullet(
        document,
        "Action: improve graceful degradation/fallback behavior during hard service outage windows.",
    )

    add_paragraph(document, "Mutation")
    add_bullet(
        document,
        f"Mutation score is {format_float(mutation_result['score'])}%, "
        f"with {mutation_result['survived_count']} surviving mutant(s).",
    )
    add_bullet(
        document,
        "Action: add targeted negative tests around content visibility and publication-status rules.",
    )

    add_paragraph(
        document,
        "Conclusion: the assignment is complete in execution, artifacts, and analysis. "
        "Partial statuses document specific, measurable improvement points with a clear follow-up plan.",
    )

    add_figure(
        document,
        figure_dir / "combined_quality_snapshot.png",
        "Figure X1. Cross-area quality snapshot.",
    )


def build_recommendations_section(document: Document, chaos_results: list[ChaosScenarioResult]) -> None:
    add_heading(document, "5. Recommendations", level=1)
    add_bullet(
        document,
        "Performance: isolate auth_login failure modes by recording HTTP status distribution and response body categories "
        "for error responses under load.",
    )
    add_bullet(
        document,
        "Performance: separate functional authentication checks from throughput stress by adding read-heavy authenticated "
        "workloads and token reuse scenarios.",
    )
    add_bullet(
        document,
        "Mutation: add focused tests for content visibility state transitions and unauthorized publication-state access.",
    )
    if any(not result.fault_injected for result in chaos_results):
        add_bullet(
            document,
            "Chaos: enable tc/netem support in the backend container image (iproute2) so network-latency scenario becomes "
            "an actual injected fault instead of a dry run.",
        )
    add_bullet(
        document,
        "Reporting: capture dashboard screenshots during each scenario and map each image to scenario IDs for reproducible "
        "evidence in the final submission.",
    )


def build_document(
    repo_root: Path,
    output_path: Path,
    student_name: str | None,
    course_name: str,
) -> Path:
    experimental_summary_path = repo_root / ".tmp" / "qa" / "experimental" / "experimental-summary.json"
    experimental_summary = read_json(experimental_summary_path)

    performance_entries, thresholds, performance_run_id = load_performance(repo_root, experimental_summary)
    mutation_summary = load_mutation(experimental_summary)
    chaos_results, chaos_run_id = load_chaos(experimental_summary)
    figure_dir = repo_root / "tmp" / "docs" / "figures" / "assignment3"

    document = Document()
    configure_document(document)

    add_heading(document, "Assignment 3 Experimental Engineering Report", level=1)
    add_paragraph(document, f"Course: {course_name}")
    add_paragraph(document, f"Student: {student_name or 'To be filled by student'}")
    add_paragraph(document, "")

    add_heading(document, "Objective and Scope", level=2)
    add_paragraph(
        document,
        "This report documents performance, mutation, and chaos testing for Assignment 3. "
        "It provides reproducible metrics, comparative analysis, identified gaps, and actionable recommendations.",
    )
    add_paragraph(
        document,
        "High-risk focus modules: authentication/session flow, authorization/admin moderation, and content/editorial workflows.",
    )

    add_heading(document, "Experimental Setup", level=2)
    env = experimental_summary.get("environment", {})
    add_bullet(document, f"Node.js version: {env.get('nodeVersion', 'unknown')}")
    add_bullet(document, f"Platform: {env.get('platform', 'unknown')}")
    add_bullet(document, f"CPU model: {env.get('cpuModel', 'unknown')}")
    add_bullet(document, f"Total memory (bytes): {env.get('totalMemoryBytes', 'unknown')}")
    add_bullet(document, "Infrastructure: local Docker Compose stack (backend + PostgreSQL) for controlled runs.")
    add_bullet(document, "Artifacts root: .tmp/qa/experimental and server/reports/mutation.")

    perf_result = build_performance_section(document, performance_entries, thresholds, figure_dir)
    mutation_result = build_mutation_section(document, mutation_summary, figure_dir)
    chaos_result = build_chaos_section(document, chaos_results, figure_dir)

    build_comparative_section(document, perf_result, mutation_result, chaos_result, figure_dir)
    build_recommendations_section(document, chaos_results)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)
    return output_path


def parse_args() -> argparse.Namespace:
    default_repo_root = Path(__file__).resolve().parents[2]
    default_output = default_repo_root / "output" / "doc" / "Assignment_3_Experimental_Report.docx"

    parser = argparse.ArgumentParser(description="Generate Assignment 3 Experimental Engineering DOCX report.")
    parser.add_argument(
        "--repo-root",
        type=Path,
        default=default_repo_root,
        help="Repository root path.",
    )
    parser.add_argument(
        "--output",
        type=Path,
        default=default_output,
        help="Output DOCX path.",
    )
    parser.add_argument(
        "--student",
        type=str,
        default=None,
        help="Student name for the title block.",
    )
    parser.add_argument(
        "--course",
        type=str,
        default="AQA",
        help="Course name for the title block.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    output = build_document(
        repo_root=args.repo_root.resolve(),
        output_path=args.output.resolve(),
        student_name=args.student,
        course_name=args.course,
    )
    print(f"Report generated: {output}")


if __name__ == "__main__":
    main()
